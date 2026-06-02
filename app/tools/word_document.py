from __future__ import annotations

import os
from typing import Any, Type

from docx import Document
from pydantic import BaseModel, Field
from langchain_core.tools import BaseTool

from app.utils import ResourceUtils


class WordDocumentInput(BaseModel):
    file_name: str = Field(description="生成的 Word 文件名，支持传入 .docx 或不带后缀的名称")
    title: str = Field(description="文档标题")
    content: str = Field(description="文档正文，支持使用换行分段")


class WordDocumentTool(BaseTool):
    """生成 Word 文档工具。"""

    name: str = "word_document_tool"
    description: str = "当需要把内容整理成 Word 文档并保存到本地时使用该工具"
    args_schema: Type[BaseModel] = WordDocumentInput

    def _run(self, *args: Any, **kwargs: Any) -> str:
        file_name = self._normalize_file_name(kwargs.get("file_name"))
        title = (kwargs.get("title") or "").strip()
        content = (kwargs.get("content") or "").strip()

        if not file_name:
            return "错误：缺少文件名 file_name。"
        if not title:
            return "错误：缺少文档标题 title。"
        if not content:
            return "错误：缺少文档正文 content。"

        output_dir = ResourceUtils.get_resource_path("generated_docs")
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, file_name)

        document = Document()
        document.add_heading(title, level=1)
        for paragraph in self._split_paragraphs(content):
            document.add_paragraph(paragraph)

        document.save(output_path)
        return f"已生成 Word 文档，保存路径：{output_path}"

    @staticmethod
    def _normalize_file_name(file_name: Any) -> str:
        normalized = os.path.basename((file_name or "").strip())
        if not normalized:
            return ""
        if not normalized.lower().endswith(".docx"):
            normalized = f"{normalized}.docx"
        return normalized

    @staticmethod
    def _split_paragraphs(content: str) -> list[str]:
        paragraphs = [line.strip() for line in content.splitlines() if line.strip()]
        return paragraphs or [content.strip()]